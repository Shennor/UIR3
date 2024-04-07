#
#    Copyright 2017 Vitalii Kulanov
#

__all__ = ['DocumentWrapper']

import io
import os
import re
import zipfile

from PIL import Image, ImageOps
from PIL.PngImagePlugin import PngInfo
from docx import Document

from docx.enum.shape import WD_INLINE_SHAPE


class DocumentWrapper(object):
    """Wrapper class for retrieving docx document attributes."""

    def __init__(self, filename):
        self._grayscale_images = False
        self._filename = filename
        self._document = Document(filename)

    def iter_paragraphs(self, styles=None):
        """Get paragraphs of specific styles of a document.

        :param styles: Paragraph styles (as a list of strings) that have
                       to be fetched. None value implies all paragraphs
        :type styles: list
        """

        for paragraph in self._document.paragraphs:
            if styles:
                if paragraph.style.name in styles:
                    yield paragraph
            else:
                yield paragraph

    def iter_sections(self):
        """Iterate over sections in docx document."""

        for section in self._document.sections:
            yield section

    def iter_runs(self, paragraph=None):
        if paragraph is None:
            for i, paragraph in enumerate(self.iter_paragraphs()):
                for run in paragraph.runs:
                    yield run
        else:
            for run in paragraph.runs:
                yield run

    def get_paragraph_attributes(self, paragraph, unit='cm'):
        """Get attributes for specified paragraph."""

        _except_attributes = ('tab_stops',)

        fetched_attributes = {}
        for attr, member in type(paragraph.paragraph_format).__dict__.items():
            if isinstance(member, property) and attr not in _except_attributes:
                fetched_attributes[attr] = self._convert_unit(
                    paragraph.paragraph_format.__getattribute__(attr) or
                    self.find_paragraph_attribute(paragraph.style,
                                                  'paragraph_format',
                                                  attr),
                    unit)
        return fetched_attributes

    def find_paragraph_attribute(self, p_style, p_element, attr):
        value = p_style.__getattribute__(p_element).__getattribute__(attr)
        if value is None and p_style.base_style is not None:
            return self.find_paragraph_attribute(p_style.base_style,
                                                 p_element, attr)
        return value

    def get_word_count(self):
        cnt = 0
        for i, paragraph in enumerate(self.iter_paragraphs()):
            cnt += len(re.findall(r'\w+', paragraph.text))
        return cnt

    def get_symbol_count_with_spaces_count(self):
        cnt = 0
        for i, paragraph in enumerate(self.iter_paragraphs()):
            cnt += len(paragraph.text) - len(re.findall(r'\n', paragraph.text))
        return cnt

    def get_symbol_count_without_spaces_count(self):
        cnt = 0
        for i, paragraph in enumerate(self.iter_paragraphs()):
            cnt += len(paragraph.text) - len(re.findall(r' ', paragraph.text)) \
                   - len(re.findall(r'\t', paragraph.text)) \
                   - len(re.findall(r'\n', paragraph.text))
        return cnt

    def get_images_shapes(self):
        images = []
        for shape in self._document.inline_shapes:
            if shape.type == WD_INLINE_SHAPE.PICTURE or shape.type == WD_INLINE_SHAPE.LINKED_PICTURE \
                    or shape.type == WD_INLINE_SHAPE.CHART:
                images.append(shape)
        return images

    def get_images_files(self):
        images = []
        archive = zipfile.ZipFile(self._filename)
        for file in archive.filelist:
            if file.filename.startswith('word/media/image'):
                image = Image.open(str(archive.extract(file)))
                images.append(image)
                '''        
            for rel in self._document.part.rels.values():
            if "image" in rel.reltype:
                image = Image.open(f"{self._filename}{rel.target_part.partname}")
                images.append(image)
                '''

        return images

    def grayscale_images(self):
        """
        image_files = self.get_images_files()
        if i < 0 or i > len(image_files):
            return False
        k = 0
        for rel in self._document.part.rels.values():
            if "image" in rel.reltype:
                if k == i:
                    rel.target_part.blob = (ImageOps.grayscale(image_files[i])).tobytes()
                    return True
                else:
                    k += 1
        """
        # print(os.access(self._filename, os.W_OK))
        self._grayscale_images = True


    def get_tables(self):
        return self._document.tables

    def save_as(self, filepath):
        self._document.save(filepath)
        if self._grayscale_images:
            archive = zipfile.ZipFile(filepath, "a")
            done = []
            num = 0
            infolist = archive.infolist()
            for entry in infolist:
                if entry.filename.startswith(f'word/media/image') and entry.filename not in done:
                    file = archive.open(entry, "r")
                    image = Image.open(file)
                    image = ImageOps.grayscale(image)
                    bs = io.BytesIO()
                    image.save(bs, format=entry.filename.split(".")[-1])
                    bs = bs.getvalue()
                    file.close()
                    file = archive.open(entry, "w")
                    file.write(bs)
                    file.close()
                    done.append(entry.filename)
                    num += 1
            archive.close()
            return num


    @staticmethod
    def _convert_unit(value, unit):
        try:
            value = value.__getattribute__(unit)
        except AttributeError:
            pass
        return value
